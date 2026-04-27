"""
DocMind 文件解析器
支持 TXT / PDF / Word (.docx) / Markdown 格式的文档解析
每种解析器返回 List[Document]，携带 metadata（来源文件名、页码/段落号）
"""
import os
from typing import List

from langchain_core.documents import Document
from utils.logger import get_logger

logger = get_logger()


def parse_txt(file_path: str) -> List[Document]:
    """
    解析 TXT 文件
    
    Args:
        file_path: 文件路径
    
    Returns:
        Document 列表
    """
    logger.info(f"解析 TXT 文件: {file_path}")
    
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    doc = Document(
        page_content=content,
        metadata={
            "source": os.path.basename(file_path),
            "file_type": "txt",
        }
    )
    
    logger.info(f"TXT 解析完成，共 {len(content)} 字符")
    return [doc]


def parse_pdf(file_path: str) -> List[Document]:
    """
    解析 PDF 文件
    Week 2 优化：使用 pdfplumber 提取表格转换为 Markdown，结合 PyMuPDF 提取纯文本。
    
    Args:
        file_path: 文件路径
    
    Returns:
        Document 列表，每页一个 Document
    """
    logger.info(f"解析 PDF 文件: {file_path}")
    
    try:
        import fitz  # PyMuPDF
        import pdfplumber
    except ImportError:
        raise ImportError("请安装 PyMuPDF 和 pdfplumber: pip install PyMuPDF pdfplumber")
    
    docs = []
    
    try:
        pdf_doc = fitz.open(file_path)
        total_pages = len(pdf_doc)
        
        with pdfplumber.open(file_path) as plumber_pdf:
            for page_num in range(total_pages):
                # 1. 使用 PyMuPDF 提取基础文本
                page_fitz = pdf_doc[page_num]
                text = page_fitz.get_text("text")
                
                # 2. 使用 pdfplumber 提取表格
                page_plumber = plumber_pdf.pages[page_num]
                tables = page_plumber.extract_tables()
                
                table_md = ""
                if tables:
                    for table in tables:
                        # 过滤空行并转为 Markdown 格式
                        valid_rows = [[str(cell).replace("\n", " ") if cell else "" for cell in row] for row in table if any(row)]
                        if not valid_rows:
                            continue
                        
                        # 构建 Markdown 表格
                        table_md += "\n\n"
                        # 表头
                        header = valid_rows[0]
                        table_md += "| " + " | ".join(header) + " |\n"
                        # 分割线
                        table_md += "|---" * len(header) + "|\n"
                        # 数据行
                        for row in valid_rows[1:]:
                            # 补齐列数
                            if len(row) < len(header):
                                row.extend([""] * (len(header) - len(row)))
                            elif len(row) > len(header):
                                row = row[:len(header)]
                            table_md += "| " + " | ".join(row) + " |\n"
                
                # 合并文本和表格
                full_content = text
                if table_md:
                    full_content += "\n\n[表格数据]:" + table_md
                
                if full_content.strip():  # 跳过完全空白页
                    doc = Document(
                        page_content=full_content,
                        metadata={
                            "source": os.path.basename(file_path),
                            "file_type": "pdf",
                            "page": page_num + 1,
                            "total_pages": total_pages,
                        }
                    )
                    docs.append(doc)
    finally:
        if 'pdf_doc' in locals():
            pdf_doc.close()
            
    logger.info(f"PDF 解析完成，共 {total_pages} 页，有效页 {len(docs)} 页，含表格支持")
    return docs


def parse_docx(file_path: str) -> List[Document]:
    """
    解析 Word (.docx) 文件
    使用 python-docx 提取段落文本
    
    Args:
        file_path: 文件路径
    
    Returns:
        Document 列表
    """
    logger.info(f"解析 Word 文件: {file_path}")
    
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")
    
    docx_file = DocxDocument(file_path)
    
    # 将所有非空段落合并，同时记录段落编号
    docs = []
    paragraphs = []
    
    for i, para in enumerate(docx_file.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    if paragraphs:
        # 将全部段落合并为一个 Document（后续由 text_splitter 切分）
        full_text = "\n".join(paragraphs)
        doc = Document(
            page_content=full_text,
            metadata={
                "source": os.path.basename(file_path),
                "file_type": "docx",
                "total_paragraphs": len(paragraphs),
            }
        )
        docs.append(doc)
    
    logger.info(f"Word 解析完成，共 {len(paragraphs)} 个有效段落")
    return docs


def parse_markdown(file_path: str) -> List[Document]:
    """
    解析 Markdown 文件
    直接按文本加载（Markdown 本身就是纯文本）
    
    Args:
        file_path: 文件路径
    
    Returns:
        Document 列表
    """
    logger.info(f"解析 Markdown 文件: {file_path}")
    
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    doc = Document(
        page_content=content,
        metadata={
            "source": os.path.basename(file_path),
            "file_type": "markdown",
        }
    )
    
    logger.info(f"Markdown 解析完成，共 {len(content)} 字符")
    return [doc]


# 支持的文件扩展名 -> 解析函数 映射表
_PARSER_MAP = {
    ".txt": parse_txt,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".md": parse_markdown,
}

# 对外暴露支持的文件类型列表
SUPPORTED_EXTENSIONS = list(_PARSER_MAP.keys())


def parse_file(file_path: str) -> List[Document]:
    """
    根据文件后缀自动选择解析器
    
    Args:
        file_path: 文件路径
    
    Returns:
        Document 列表
    
    Raises:
        ValueError: 不支持的文件格式
        FileNotFoundError: 文件不存在
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    parser = _PARSER_MAP.get(ext)
    if parser is None:
        supported = ", ".join(SUPPORTED_EXTENSIONS)
        raise ValueError(f"不支持的文件格式: {ext}。支持的格式: {supported}")
    
    return parser(file_path)
